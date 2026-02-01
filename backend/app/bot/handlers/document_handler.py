"""
Обработчики загрузки документов в Telegram бот
Поддерживает:
- Загрузку PDF, Excel, Word, TXT файлов
- Индексацию документов в Qdrant для RAG
"""
from aiogram import Dispatcher, F
from aiogram.types import Message, Document as TelegramDocument
from aiogram.enums import ChatAction
from aiogram.fsm.context import FSMContext
from uuid import UUID
import logging
import os
import tempfile
import asyncio
from pathlib import Path
import uuid as uuid_module
from datetime import datetime

from app.core.database import AsyncSessionLocal
from app.bot.handlers.auth_handler import AuthStates
from app.models.document import Document
from app.tasks.document_tasks import process_document_task, process_large_document_with_langgraph, process_document_async
from app.services.document_agent_adapter import DocumentAgentAdapter
from app.core.config import settings

logger = logging.getLogger(__name__)


async def extract_text_from_file(file_path: str, file_extension: str) -> str:
    """Извлечение текста из различных форматов файлов"""
    try:
        if file_extension == 'pdf':
            # PDF
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            return text
        
        elif file_extension == 'docx':
            # Word документы
            try:
                import docx
                doc = docx.Document(file_path)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
                return text
            except ImportError:
                logger.warning("python-docx not installed, trying alternative method")
                return ""
        
        elif file_extension in ['xlsx', 'xls']:
            # Excel файлы
            try:
                import pandas as pd
                df = pd.read_excel(file_path)
                return df.to_string()
            except ImportError:
                logger.warning("pandas/openpyxl not installed")
                return ""
        
        elif file_extension == 'txt' or file_extension == 'md':
            # Текстовые файлы
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        else:
            logger.warning(f"Unsupported file format: {file_extension}")
            return ""
            
    except Exception as e:
        logger.error(f"Error extracting text from {file_extension} file: {e}")
        return ""


async def index_document_to_qdrant(
    text_content: str,
    file_name: str,
    user_id: str,
    username: str,
    project_id: str = None
) -> dict:
    """
    Индексация документа в Qdrant (коллекция 'data')
    
    Args:
        text_content: Текст документа
        file_name: Имя файла
        user_id: ID пользователя Telegram
        username: Username пользователя
        project_id: ID проекта (опционально)
    
    Returns:
        {"success": bool, "chunks_count": int, "error": str}
    """
    try:
        from app.services.rag.qdrant_helper import index_document_chunks_to_qdrant
        
        # Создаем уникальный ID для документа
        doc_id = str(uuid_module.uuid4())
        
        # Разбиваем текст на чанки
        try:
            from langchain_text_splitters import RecursiveCharacterTextSplitter
        except ImportError:
            # Fallback на простой сплиттер
            class RecursiveCharacterTextSplitter:
                def __init__(self, chunk_size=500, chunk_overlap=50, separators=None):
                    self.chunk_size = chunk_size
                    self.chunk_overlap = chunk_overlap
                    self.separators = separators or ["\n\n", "\n", ". ", " ", ""]
                
                def split_text(self, text):
                    chunks = []
                    current_chunk = ""
                    
                    for sep in self.separators:
                        if sep in text:
                            parts = text.split(sep)
                            for part in parts:
                                if len(current_chunk) + len(part) + len(sep) <= self.chunk_size:
                                    current_chunk += part + sep
                                else:
                                    if current_chunk:
                                        chunks.append(current_chunk.strip())
                                    current_chunk = part + sep
                            break
                    
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    
                    # Если нет чанков, разбиваем по размеру
                    if not chunks:
                        for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
                            chunk = text[i:i + self.chunk_size]
                            if chunk.strip():
                                chunks.append(chunk.strip())
                    
                    return chunks
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        chunks = text_splitter.split_text(text_content)
        
        # Фильтруем короткие чанки
        chunks = [chunk for chunk in chunks if len(chunk.strip()) >= 10]
        
        logger.info(f"📄 Создано {len(chunks)} чанков из документа {file_name}")
        
        if not chunks:
            return {
                "success": False,
                "error": "Не удалось извлечь текст из документа"
            }
        
        # Индексируем чанки в Qdrant
        result = await index_document_chunks_to_qdrant(
            chunks=chunks,
            file_name=file_name,
            doc_id=doc_id,
            user_id=user_id,
            username=username,
            project_id=project_id
        )
        
        return result
        
    except Exception as e:
        logger.error(f"❌ Ошибка индексации документа в Qdrant: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return {
            "success": False,
            "error": str(e)
        }


async def handle_document(message: Message, state: FSMContext):
    """Обработка загрузки документа (PDF, Excel)"""
    # Проверка авторизации
    current_state = await state.get_state()
    if current_state != AuthStates.authorized:
        await message.answer("❌ Сначала авторизуйтесь через /start")
        return
    
    # Получаем project_id из состояния
    data = await state.get_data()
    project_id_str = data.get("project_id")
    
    if not project_id_str:
        await message.answer("❌ Ошибка: проект не найден. Используйте /start")
        return
    
    try:
        project_id = UUID(project_id_str)
    except ValueError:
        await message.answer("❌ Ошибка: неверный ID проекта")
        return
    
    # Проверяем, что это документ
    if not message.document:
        await message.answer("❌ Пожалуйста, отправьте файл (PDF или Excel)")
        return
    
    doc = message.document
    file_name = doc.file_name or "document"
    file_size = doc.file_size or 0
    
    # Проверяем размер файла (максимум 50MB)
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    if file_size > MAX_FILE_SIZE:
        await message.answer(f"❌ Файл слишком большой. Максимальный размер: {MAX_FILE_SIZE / 1024 / 1024:.0f}MB")
        return
    
    # Определяем тип файла
    file_ext = Path(file_name).suffix.lower()
    if file_ext not in ['.pdf', '.xlsx', '.xls', '.docx', '.txt']:
        await message.answer("❌ Поддерживаются только файлы: PDF, Excel (.xlsx, .xls), Word (.docx), TXT")
        return
    
    file_type = file_ext.lstrip('.')
    
    # Показываем индикатор печати
    try:
        await message.bot.send_chat_action(message.chat.id, ChatAction.TYPING)
    except Exception as e:
        logger.warning(f"Failed to send typing indicator: {e}")
    
    # Отправляем сообщение о начале обработки
    processing_msg = await message.answer(f"📥 Загружаю файл: {file_name}...")
    
    try:
        # Скачиваем файл
        bot = message.bot
        file_info = await bot.get_file(doc.file_id)
        
        # Создаем временный файл
        temp_dir = Path(tempfile.gettempdir())
        temp_file = temp_dir / f"telegram_upload_{doc.file_id}_{file_name}"
        
        # Скачиваем файл
        await bot.download_file(file_info.file_path, destination=temp_file)
        
        logger.info(f"[TELEGRAM UPLOAD] File downloaded: {temp_file}, size: {file_size} bytes")
        
        # Читаем файл в память для обработки
        with open(temp_file, 'rb') as f:
            file_content = f.read()
        
        # Удаляем временный файл
        try:
            os.unlink(temp_file)
        except Exception as e:
            logger.warning(f"Failed to delete temp file: {e}")
        
        async with AsyncSessionLocal() as db:
            # Создаем документ в БД
            # Проверяем наличие колонки fast_mode перед созданием документа
            from sqlalchemy import inspect, text
            try:
                # Для async сессии нужно использовать другой способ проверки
                result = await db.execute(text("""
                    SELECT column_name 
                    FROM information_schema.columns 
                    WHERE table_name = 'documents' AND column_name = 'fast_mode'
                """))
                has_fast_mode = result.fetchone() is not None
            except Exception:
                # Если не удалось проверить (например, SQLite), предполагаем что колонка есть
                has_fast_mode = True
            
            document = Document(
                project_id=project_id,
                filename=file_name,
                content="Обработка...",  # Временный placeholder
                file_type=file_type
            )
            
            # Устанавливаем fast_mode только если колонка существует
            if has_fast_mode:
                document.fast_mode = False
                db.add(document)
                await db.commit()
                await db.refresh(document)
            else:
                # Если колонки нет, используем raw SQL для вставки без fast_mode
                logger.warning("[TELEGRAM UPLOAD] Колонка fast_mode не найдена, используем raw SQL")
                document_id = uuid_module.uuid4()
                await db.execute(text("""
                    INSERT INTO documents (id, project_id, filename, content, file_type, summary, created_at)
                    VALUES (:id, :project_id, :filename, :content, :file_type, :summary, :created_at)
                """), {
                    "id": str(document_id),
                    "project_id": str(project_id),
                    "filename": file_name,
                    "content": "Обработка...",
                    "file_type": file_type,
                    "summary": None,
                    "created_at": datetime.utcnow()
                })
                await db.commit()
                # Получаем созданный документ через ORM
                document = await db.get(Document, document_id)
                if not document:
                    # Если не получилось через ORM, создаем объект вручную
                    result = await db.execute(
                        text("SELECT id, project_id, filename, content, file_type, summary, created_at FROM documents WHERE id = :id"),
                        {"id": str(document_id)}
                    )
                    row = result.fetchone()
                    if row:
                        document = Document(
                            id=UUID(row[0]) if isinstance(row[0], str) else row[0],
                            project_id=UUID(row[1]) if isinstance(row[1], str) else row[1],
                            filename=row[2],
                            content=row[3],
                            file_type=row[4],
                            summary=row[5],
                            created_at=row[6] if row[6] else datetime.utcnow()
                        )
            
            logger.info(f"[TELEGRAM UPLOAD] Document created in DB: {document.id}")
            
            # Сохраняем файл в постоянное место для возможности скачивания
            media_dir = Path("media") / "documents" / str(project_id)
            media_dir.mkdir(parents=True, exist_ok=True)
            
            # Сохраняем оригинальный файл для скачивания
            original_file_path = media_dir / f"{document.id}_{file_name}"
            with open(original_file_path, 'wb') as f:
                f.write(file_content)
            logger.info(f"[TELEGRAM UPLOAD] Original file saved: {original_file_path}")
            
            # Определяем размер файла для выбора стратегии обработки
            # Для небольших файлов (< 1MB) используем синхронную обработку без Celery
            SMALL_FILE_THRESHOLD = 1 * 1024 * 1024  # 1MB
            LARGE_PDF_THRESHOLD = 5 * 1024 * 1024  # 5MB для больших PDF
            
            is_small_file = file_size < SMALL_FILE_THRESHOLD
            is_large_pdf = (
                file_type == "pdf" and 
                file_size > LARGE_PDF_THRESHOLD
            )
            
            # Быстрая проверка размера для PDF
            if is_large_pdf:
                try:
                    adapter = DocumentAgentAdapter()
                    preview_text = await adapter._quick_pdf_preview(file_content)
                    estimated_pages = len(preview_text) // 3000 if preview_text else 0
                    
                    if estimated_pages > 100:
                        is_large_pdf = True
                        logger.info(f"[TELEGRAM UPLOAD] Большой PDF обнаружен: ~{estimated_pages} страниц, используем быструю индексацию")
                    else:
                        is_large_pdf = False
                except Exception as e:
                    logger.warning(f"[TELEGRAM UPLOAD] Не удалось оценить размер PDF: {e}")
                    is_large_pdf = False
            
            # Проверяем доступность Celery
            celery_available = bool(settings.CELERY_BROKER_URL and settings.CELERY_RESULT_BACKEND)
            
            # Выбираем стратегию обработки
            if is_small_file or not celery_available:
                # Для небольших файлов или если Celery недоступен - извлекаем текст СРАЗУ и сохраняем в БД
                if not celery_available:
                    logger.warning(f"[TELEGRAM UPLOAD] Celery недоступен (Redis не настроен), извлекаем текст сразу")
                else:
                    logger.info(f"[TELEGRAM UPLOAD] Небольшой файл ({file_size / 1024:.1f} KB), извлекаем текст сразу и сохраняем в БД")
                
                # Для маленьких файлов извлекаем текст СРАЗУ и сохраняем в БД
                try:
                    # Показываем typing indicator во время извлечения текста
                    try:
                        await message.bot.send_chat_action(message.chat.id, ChatAction.TYPING)
                    except:
                        pass
                    
                    # Создаем временный файл для извлечения текста
                    temp_extract_path = temp_dir / f"extract_immediate_{document.id}_{file_name}"
                    with open(temp_extract_path, 'wb') as f:
                        f.write(file_content)
                    
                    try:
                        # Извлекаем текст из файла
                        logger.info(f"[TELEGRAM UPLOAD] Извлечение текста из файла {file_name}...")
                        text_content = await extract_text_from_file(str(temp_extract_path), file_type)
                        
                        if text_content and text_content.strip():
                            # Сохраняем контент в БД СРАЗУ
                            MAX_CONTENT_SIZE = 2_000_000
                            if len(text_content) > MAX_CONTENT_SIZE:
                                logger.warning(f"[TELEGRAM UPLOAD] Текст слишком большой ({len(text_content)} символов), обрезаем до {MAX_CONTENT_SIZE}")
                                document.content = text_content[:MAX_CONTENT_SIZE] + f"\n\n[... документ обрезан, всего {len(text_content)} символов ...]"
                            else:
                                document.content = text_content
                            
                            await db.commit()
                            await db.refresh(document)
                            
                            saved_length = len(document.content) if document.content else 0
                            logger.info(f"[TELEGRAM UPLOAD] ✅ Текст извлечен и сохранен в БД: {saved_length} символов для документа {document.id}")
                        else:
                            logger.warning(f"[TELEGRAM UPLOAD] ⚠️ Не удалось извлечь текст из файла {file_name}")
                            document.content = "Ошибка извлечения текста"
                            await db.commit()
                    finally:
                        # Удаляем временный файл
                        try:
                            os.unlink(temp_extract_path)
                        except:
                            pass
                    
                    # Запускаем фоновую обработку для дополнительных задач (чанки, метаданные и т.д.)
                    asyncio.create_task(
                        process_document_async(
                            document.id,
                            project_id,
                            file_content,
                            file_name,
                            file_type
                        )
                    )
                    logger.info(f"[TELEGRAM UPLOAD] Фоновая обработка запущена для документа {document.id} (контент уже в БД)")
                except Exception as extract_error:
                    logger.error(f"[TELEGRAM UPLOAD] Ошибка извлечения текста: {extract_error}", exc_info=True)
                    # В случае ошибки запускаем обычную обработку
                    asyncio.create_task(
                        process_document_async(
                            document.id,
                            project_id,
                            file_content,
                            file_name,
                            file_type
                        )
                    )
            else:
                # Для больших файлов используем Celery
                # Сохраняем файл во временное место для Celery задачи
                temp_path = temp_dir / f"celery_doc_{document.id}_{file_name}"
                with open(temp_path, 'wb') as f:
                    f.write(file_content)
                
                if is_large_pdf:
                    # Используем оптимизированную обработку для больших PDF
                    logger.info(f"[TELEGRAM UPLOAD] Используем быструю индексацию для большого PDF через Celery")
                    try:
                        task_result = process_large_document_with_langgraph.delay(
                            str(document.id),
                            str(project_id),
                            str(temp_path),
                            file_name,
                            file_type
                        )
                        logger.info(f"[TELEGRAM UPLOAD] Celery task created: {task_result.id} for document {document.id}, is_large_pdf: {is_large_pdf}")
                    except Exception as celery_error:
                        logger.error(f"[TELEGRAM UPLOAD] Ошибка создания Celery задачи: {celery_error}, используем синхронную обработку")
                        # Fallback на синхронную обработку
                        asyncio.create_task(
                            process_document_async(
                                document.id,
                                project_id,
                                file_content,
                                file_name,
                                file_type
                            )
                        )
                else:
                    # Обычная обработка через Celery
                    try:
                        task_result = process_document_task.delay(
                            str(document.id),
                            str(project_id),
                            str(temp_path),
                            file_name,
                            file_type
                        )
                        logger.info(f"[TELEGRAM UPLOAD] Celery task created: {task_result.id} for document {document.id}")
                    except Exception as celery_error:
                        logger.error(f"[TELEGRAM UPLOAD] Ошибка создания Celery задачи: {celery_error}, используем синхронную обработку")
                        # Fallback на синхронную обработку
                        asyncio.create_task(
                            process_document_async(
                                document.id,
                                project_id,
                                file_content,
                                file_name,
                                file_type
                            )
                        )
            
            # Для коротких файлов НЕ индексируем в Qdrant - используем только контент из БД
            # RAG только для больших файлов и фоновой обработки
            if is_small_file:
                # Короткий файл - просто сохраняем, контент будет использован напрямую из БД
                logger.info(f"[TELEGRAM UPLOAD] Короткий файл ({file_size / 1024:.1f} KB), пропускаем индексацию в Qdrant, используем контент из БД")
                
                # Показываем typing перед отправкой ответа
                try:
                    await message.bot.send_chat_action(message.chat.id, ChatAction.TYPING)
                except:
                    pass
                
                status_text = (
                    f"✅ <b>Файл успешно загружен!</b>\n\n"
                    f"📄 Название: {file_name}\n"
                    f"📊 Тип: {file_type.upper()}\n"
                    f"📏 Размер: {file_size / 1024:.1f} KB\n\n"
                    f"💡 Для коротких файлов используется прямой доступ к содержимому.\n"
                    f"📚 Документ готов для вопросов!\n"
                    f"Используйте /documents для просмотра списка документов."
                )
                await processing_msg.edit_text(status_text, parse_mode="HTML")
                
                # Сохраняем document_id в state для приоритета при поиске
                await state.update_data(last_document_id=str(document.id))
                logger.info(f"[TELEGRAM UPLOAD] Saved last_document_id={document.id} to state (small file, no Qdrant)")
            else:
                # Для больших файлов индексируем в Qdrant для RAG
                try:
                    # Показываем typing indicator во время обработки
                    try:
                        await message.bot.send_chat_action(message.chat.id, ChatAction.TYPING)
                    except:
                        pass
                    
                    # Извлекаем текст из файла для Qdrant
                    # Для больших файлов используем temp_path
                    text_content = await extract_text_from_file(str(temp_path), file_type)
                    
                    if text_content and text_content.strip():
                        # Получаем данные пользователя
                        telegram_user_id = str(message.from_user.id)
                        telegram_username = message.from_user.username or "unknown"
                        
                        # Показываем typing indicator перед индексацией
                        try:
                            await message.bot.send_chat_action(message.chat.id, ChatAction.TYPING)
                        except:
                            pass
                        
                        # Индексируем в Qdrant
                        qdrant_result = await index_document_to_qdrant(
                            text_content=text_content,
                            file_name=file_name,
                            user_id=telegram_user_id,
                            username=telegram_username,
                            project_id=str(project_id)
                        )
                        
                        if qdrant_result.get("success"):
                            chunks_count = qdrant_result.get("chunks_count", 0)
                            logger.info(f"[TELEGRAM UPLOAD] ✅ Document indexed in Qdrant: {chunks_count} chunks")
                            
                            # Показываем typing перед отправкой ответа
                            try:
                                await message.bot.send_chat_action(message.chat.id, ChatAction.TYPING)
                            except:
                                pass
                            
                            status_text = (
                                f"✅ <b>Файл успешно загружен и обработан!</b>\n\n"
                                f"📄 Название: {file_name}\n"
                                f"📊 Тип: {file_type.upper()}\n"
                                f"📏 Размер: {file_size / 1024 / 1024:.2f} MB\n"
                                f"🔍 Чанков в RAG: {chunks_count}\n"
                            )
                            if is_large_pdf:
                                status_text += f"⚡ Используется быстрая индексация для большого PDF\n"
                            status_text += (
                                f"\n⏳ Полная обработка документа продолжается в фоне.\n"
                                f"📚 Документ уже доступен для поиска!\n"
                                f"Используйте /documents для просмотра списка документов."
                            )
                            await processing_msg.edit_text(status_text, parse_mode="HTML")
                            
                            # Сохраняем document_id в state для приоритета при поиске
                            await state.update_data(last_document_id=str(document.id))
                            logger.info(f"[TELEGRAM UPLOAD] Saved last_document_id={document.id} to state")
                        else:
                            error_msg = qdrant_result.get("error", "Unknown error")
                            logger.warning(f"[TELEGRAM UPLOAD] ⚠️ Qdrant indexing failed: {error_msg}")
                            
                            await processing_msg.edit_text(
                                f"✅ <b>Файл успешно загружен!</b>\n\n"
                                f"📄 Название: {file_name}\n"
                                f"📊 Тип: {file_type.upper()}\n"
                                f"📏 Размер: {file_size / 1024 / 1024:.2f} MB\n\n"
                                f"⏳ Обработка документа начата.\n"
                                f"⚠️ Индексация в RAG будет выполнена позже.\n"
                                f"Используйте /documents для просмотра списка документов.",
                                parse_mode="HTML"
                            )
                            
                            # Сохраняем document_id в state для приоритета при поиске
                            await state.update_data(last_document_id=str(document.id))
                            logger.info(f"[TELEGRAM UPLOAD] Saved last_document_id={document.id} to state")
                    else:
                        logger.warning(f"[TELEGRAM UPLOAD] ⚠️ No text extracted from document for Qdrant")
                        await processing_msg.edit_text(
                            f"✅ <b>Файл успешно загружен!</b>\n\n"
                            f"📄 Название: {file_name}\n"
                            f"📊 Тип: {file_type.upper()}\n"
                            f"📏 Размер: {file_size / 1024 / 1024:.2f} MB\n\n"
                            f"⏳ Обработка документа начата. Это может занять некоторое время.\n"
                            f"Используйте /documents для просмотра списка документов.",
                            parse_mode="HTML"
                        )
                        
                        # Сохраняем document_id в state для приоритета при поиске
                        await state.update_data(last_document_id=str(document.id))
                        logger.info(f"[TELEGRAM UPLOAD] Saved last_document_id={document.id} to state")
                        
                except Exception as qdrant_error:
                    logger.error(f"[TELEGRAM UPLOAD] ❌ Qdrant indexing error: {qdrant_error}")
                    import traceback
                    logger.error(traceback.format_exc())
                    
                    # Всё равно показываем успешную загрузку (Celery обработает позже)
                    await processing_msg.edit_text(
                        f"✅ <b>Файл успешно загружен!</b>\n\n"
                        f"📄 Название: {file_name}\n"
                        f"📊 Тип: {file_type.upper()}\n"
                        f"📏 Размер: {file_size / 1024 / 1024:.2f} MB\n\n"
                        f"⏳ Обработка документа начата. Это может занять некоторое время.\n"
                        f"Используйте /documents для просмотра списка документов.",
                        parse_mode="HTML"
                    )
                    
                    # Сохраняем document_id в state для приоритета при поиске
                    await state.update_data(last_document_id=str(document.id))
                    logger.info(f"[TELEGRAM UPLOAD] Saved last_document_id={document.id} to state")
    
    except Exception as e:
        logger.error(f"[TELEGRAM UPLOAD] Error uploading document: {e}", exc_info=True)
        error_msg = str(e)
        # Экранируем HTML специальные символы
        error_msg = error_msg.replace("<", "&lt;").replace(">", "&gt;")
        try:
            await processing_msg.edit_text(
                f"❌ Ошибка при загрузке файла: {error_msg[:500]}\n"
                f"Попробуйте позже или обратитесь к администратору."
            )
        except Exception as edit_error:
            # Если не удалось отредактировать, отправляем новое сообщение
            logger.warning(f"[TELEGRAM UPLOAD] Failed to edit message: {edit_error}")
            try:
                await message.answer(
                    f"❌ Ошибка при загрузке файла.\n"
                    f"Попробуйте позже или обратитесь к администратору."
                )
            except:
                pass


def register_document_handlers(dp: Dispatcher, project_id: str):
    """Регистрация обработчиков документов"""
    import logging
    logger = logging.getLogger(__name__)
    
    # Регистрируем обработчик для документов (только для авторизованных пользователей)
    dp.message.register(handle_document, AuthStates.authorized, F.document)
    logger.info(f"[REGISTER HANDLERS] Document handlers registered for project {project_id}")

